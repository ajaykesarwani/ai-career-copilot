"""
Document Export
================
Renders AI-generated resume/cover-letter text into properly formatted
DOCX and PDF files.

Layout awareness: when a ResumeLayout dict is provided (extracted from the
candidate's uploaded resume via layout_analyser.py), the renderer adapts:
  - accent colour matches the original
  - section order follows the original
  - two-column / sidebar layout is reproduced where possible
  - header style (centred vs left-aligned vs banner) is matched
"""

from __future__ import annotations
import io
import re
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.units import inch, cm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, HRFlowable,
    Table, TableStyle, KeepTogether,
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT

# ── Default brand colours (overridden by layout accent_color) ─────────────────
DEFAULT_BRAND   = "6C5CE7"
TEXT_HEX        = "1A1A2E"
MUTED_HEX       = "4A4A6A"
LIGHT_BG        = "F4F3FF"   # very light tint for sidebar/banner


# ── Section parsing ────────────────────────────────────────────────────────────

_KNOWN_SECTIONS = [
    "CONTACT", "SUMMARY", "OBJECTIVE", "SKILLS", "EXPERIENCE",
    "PROJECTS", "EDUCATION", "CERTIFICATIONS", "AWARDS",
    "LANGUAGES", "INTERESTS", "PUBLICATIONS",
]

def _parse_resume_sections(text: str) -> dict[str, str]:
    loose = r"(?im)^\s*(?:WORK\s+|PROFESSIONAL\s+|KEY\s+)?(" + "|".join(_KNOWN_SECTIONS) + r")\s*:?\s*$"
    matches = list(re.finditer(loose, text))
    sections: dict[str, str] = {}
    if not matches:
        sections["SUMMARY"] = text.strip()
        return sections
    preamble = text[: matches[0].start()].strip()
    if preamble:
        sections["CONTACT"] = preamble
    for i, m in enumerate(matches):
        key = m.group(1).upper()
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        body = text[start:end].strip()
        if body:
            sections[key] = body
    return sections

def _bullets(block: str) -> list[str]:
    lines = [l.strip() for l in block.split("\n") if l.strip()]
    return [re.sub(r"^[•\-\*✓▸→]\s*", "", l) for l in lines]

def _section_order(layout: Optional[dict]) -> list[str]:
    if layout and layout.get("section_order"):
        known = set(_KNOWN_SECTIONS)
        ordered = [s.upper() for s in layout["section_order"] if s.upper() in known]
        # Append any known section not already listed
        ordered += [s for s in _KNOWN_SECTIONS if s not in ordered]
        return ordered
    return _KNOWN_SECTIONS

def _accent_rgb(layout: Optional[dict]) -> tuple[int, int, int]:
    """Return (R, G, B) from layout accent_color hex, defaulting to brand purple."""
    hex_str = (layout or {}).get("accent_color", DEFAULT_BRAND) or DEFAULT_BRAND
    hex_str = hex_str.strip("#").upper()
    if len(hex_str) != 6:
        hex_str = DEFAULT_BRAND
    try:
        r = int(hex_str[0:2], 16)
        g = int(hex_str[2:4], 16)
        b = int(hex_str[4:6], 16)
        return r, g, b
    except ValueError:
        return 0x6C, 0x5C, 0xE7

def _accent_hex(layout: Optional[dict]) -> str:
    h = (layout or {}).get("accent_color", DEFAULT_BRAND) or DEFAULT_BRAND
    return h.strip("#").upper()[:6] or DEFAULT_BRAND


# ══════════════════════════════════════════════════════════════════════════════
# DOCX RENDERING
# ══════════════════════════════════════════════════════════════════════════════

def render_resume_docx(text: str, profile: dict, layout: Optional[dict] = None) -> bytes:
    sections      = _parse_resume_sections(text)
    order         = _section_order(layout)
    accent_rgb    = _accent_rgb(layout)
    header_style  = (layout or {}).get("header_style", "centered")
    uses_dividers = (layout or {}).get("uses_dividers", True)
    is_two_col    = (layout or {}).get("columns", 1) == 2 or (layout or {}).get("has_sidebar", False)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    for section in doc.sections:
        section.top_margin    = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin   = Inches(0.75)
        section.right_margin  = Inches(0.75)

    name  = profile.get("name") or "[Your Full Name]"
    title = profile.get("title") or ""

    # ── Header block ──────────────────────────────────────────────────────────
    if header_style == "banner":
        _docx_banner_header(doc, name, title, profile, accent_rgb)
    elif header_style == "left-aligned":
        _docx_left_header(doc, name, title, profile, accent_rgb)
    else:
        _docx_centered_header(doc, name, title, profile, accent_rgb)

    if uses_dividers:
        _add_hr(doc, accent_rgb)

    # ── Two-column via table ───────────────────────────────────────────────────
    if is_two_col:
        _docx_two_col_body(doc, sections, order, accent_rgb, uses_dividers)
    else:
        _docx_single_col_body(doc, sections, order, accent_rgb, uses_dividers)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render_cover_letter_docx(text: str, profile: dict, job: Optional[dict] = None) -> bytes:
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(1)
        section.left_margin = section.right_margin = Inches(1)

    name = profile.get("name") or "[Your Full Name]"
    today = datetime.now().strftime("%B %d, %Y")

    for line in [name, profile.get("email") or "[Your Email Address]",
                 profile.get("phone") or "[Your Phone Number]",
                 profile.get("location") or "[Your City, Country]"]:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(0)
        p.add_run(line).font.size = Pt(10)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    doc.add_paragraph(today)
    doc.add_paragraph()
    if job:
        for line in [job.get("company", ""), "[Company Address]"]:
            p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(0)
            p.add_run(line).font.size = Pt(10)
        doc.add_paragraph()

    body_paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text.strip()) if p.strip()]
    for para in body_paragraphs:
        sign_match = re.match(
            r"(?i)^(best regards|sincerely|kind regards|warm regards),?\s*\n?\s*(.*)$",
            para, re.DOTALL
        )
        if sign_match:
            p = doc.add_paragraph(f"{sign_match.group(1)},")
            p.paragraph_format.space_after = Pt(2)
            for r in p.runs: r.font.size = Pt(11)
            if sign_match.group(2).strip():
                p2 = doc.add_paragraph(sign_match.group(2).strip())
                p2.paragraph_format.space_after = Pt(10)
                for r in p2.runs: r.font.size = Pt(11)
            continue
        p = doc.add_paragraph(para)
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.line_spacing = 1.25
        for r in p.runs: r.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── DOCX header helpers ────────────────────────────────────────────────────────

def _contact_line(profile: dict) -> str:
    parts = [
        profile.get("email") or "[Your Email Address]",
        profile.get("phone") or "[Your Phone Number]",
        profile.get("location") or "[Your City, Country]",
        profile.get("linkedin_url") or "[Your LinkedIn URL]",
    ]
    if profile.get("github_url"): parts.append(profile["github_url"])
    return " | ".join(parts)

def _docx_centered_header(doc: Document, name: str, title: str, profile: dict, accent_rgb: tuple):
    h = doc.add_paragraph(); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run(name); r.font.size = Pt(22); r.font.bold = True
    r.font.color.rgb = RGBColor(int(TEXT_HEX[0:2],16), int(TEXT_HEX[2:4],16), int(TEXT_HEX[4:6],16))
    if title:
        t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tr = t.add_run(title); tr.font.size = Pt(12); tr.font.bold = True
        tr.font.color.rgb = RGBColor(*accent_rgb)
        t.paragraph_format.space_after = Pt(4)
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = c.add_run(_contact_line(profile))
    cr.font.size = Pt(9); cr.font.color.rgb = RGBColor(int(MUTED_HEX[0:2],16),int(MUTED_HEX[2:4],16),int(MUTED_HEX[4:6],16))
    c.paragraph_format.space_after = Pt(10)

def _docx_left_header(doc: Document, name: str, title: str, profile: dict, accent_rgb: tuple):
    h = doc.add_paragraph()
    r = h.add_run(name); r.font.size = Pt(22); r.font.bold = True
    r.font.color.rgb = RGBColor(int(TEXT_HEX[0:2],16),int(TEXT_HEX[2:4],16),int(TEXT_HEX[4:6],16))
    if title:
        t = doc.add_paragraph()
        tr = t.add_run(title); tr.font.size = Pt(12); tr.font.bold = True
        tr.font.color.rgb = RGBColor(*accent_rgb)
        t.paragraph_format.space_after = Pt(4)
    c = doc.add_paragraph()
    cr = c.add_run(_contact_line(profile))
    cr.font.size = Pt(9); cr.font.color.rgb = RGBColor(int(MUTED_HEX[0:2],16),int(MUTED_HEX[2:4],16),int(MUTED_HEX[4:6],16))
    c.paragraph_format.space_after = Pt(10)

def _docx_banner_header(doc: Document, name: str, title: str, profile: dict, accent_rgb: tuple):
    """Header with a coloured background band containing name + title."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    cell.width = Inches(7)
    # Shade the cell
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), f"{accent_rgb[0]:02X}{accent_rgb[1]:02X}{accent_rgb[2]:02X}")
    tcPr.append(shd)
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(name); r.font.size = Pt(20); r.font.bold = True; r.font.color.rgb = RGBColor(255,255,255)
    if title:
        p2 = cell.add_paragraph(title); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.runs[0].font.color.rgb = RGBColor(230,230,255); p2.runs[0].font.size = Pt(11)
    doc.add_paragraph()
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = c.add_run(_contact_line(profile))
    cr.font.size = Pt(9); cr.font.color.rgb = RGBColor(int(MUTED_HEX[0:2],16),int(MUTED_HEX[2:4],16),int(MUTED_HEX[4:6],16))
    c.paragraph_format.space_after = Pt(10)


def _docx_single_col_body(doc: Document, sections: dict, order: list, accent_rgb: tuple, uses_dividers: bool):
    labels = {"SUMMARY":"Summary","SKILLS":"Skills","EXPERIENCE":"Experience",
              "PROJECTS":"Projects","EDUCATION":"Education","CERTIFICATIONS":"Certifications",
              "AWARDS":"Awards","LANGUAGES":"Languages","INTERESTS":"Interests",
              "OBJECTIVE":"Objective","PUBLICATIONS":"Publications"}
    for key in order:
        body = sections.get(key)
        if not body: continue
        _add_section_heading_docx(doc, labels.get(key, key.title()), accent_rgb)
        if key in ("EXPERIENCE","PROJECTS"):
            for line in _bullets(body):
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(2)
                p.add_run(line).font.size = Pt(10.5)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            p.add_run(body).font.size = Pt(10.5)


def _docx_two_col_body(doc: Document, sections: dict, order: list, accent_rgb: tuple, uses_dividers: bool):
    """Render body as a two-column table: sidebar (skills/contact) | main content."""
    sidebar_keys = [k for k in ["SKILLS","CERTIFICATIONS","LANGUAGES","INTERESTS"] if sections.get(k)]
    main_keys    = [k for k in order if k not in sidebar_keys and k != "CONTACT" and sections.get(k)]

    # Build sidebar content
    left_parts  = []
    right_parts = []

    for key in sidebar_keys:
        left_parts.append((key.title(), _bullets(sections[key])))
    for key in main_keys:
        right_parts.append((key.title(), sections[key], key in ("EXPERIENCE","PROJECTS")))

    if not left_parts:
        _docx_single_col_body(doc, sections, order, accent_rgb, uses_dividers)
        return

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    left_cell  = tbl.cell(0,0)
    right_cell = tbl.cell(0,1)
    left_cell.width  = Inches(2.2)
    right_cell.width = Inches(4.8)
    # Shade left sidebar
    _shade_cell(left_cell, LIGHT_BG)

    # Populate left sidebar
    for heading, items in left_parts:
        p = left_cell.add_paragraph()
        r = p.add_run(heading.upper()); r.font.size = Pt(9); r.font.bold = True
        r.font.color.rgb = RGBColor(*accent_rgb)
        p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
        for item in items:
            bp = left_cell.add_paragraph(item)
            bp.paragraph_format.space_after = Pt(1)
            for r in bp.runs: r.font.size = Pt(9)

    # Populate right main
    for heading, body, is_bullets in right_parts:
        p = right_cell.add_paragraph()
        r = p.add_run(heading.upper()); r.font.size = Pt(10); r.font.bold = True
        r.font.color.rgb = RGBColor(*accent_rgb)
        p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
        if is_bullets:
            for line in _bullets(body):
                bp = right_cell.add_paragraph(line, style="List Bullet")
                bp.paragraph_format.space_after = Pt(1)
                for r in bp.runs: r.font.size = Pt(10.5)
        else:
            bp = right_cell.add_paragraph(body)
            bp.paragraph_format.space_after = Pt(6)
            for r in bp.runs: r.font.size = Pt(10.5)
    # Remove table borders (visual = no visible grid)
    for row in tbl.rows:
        for cell in row.cells:
            _remove_cell_borders(cell)


def _shade_cell(cell, hex_color: str):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color); tcPr.append(shd)

def _remove_cell_borders(cell):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top","left","bottom","right","insideH","insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),"none"); b.set(qn("w:sz"),"0"); b.set(qn("w:space"),"0"); b.set(qn("w:color"),"auto")
        tcBorders.append(b)
    tcPr.append(tcBorders)

def _add_hr(doc: Document, accent_rgb: tuple):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),"single"); bottom.set(qn("w:sz"),"6")
    bottom.set(qn("w:space"),"1"); bottom.set(qn("w:color"),f"{accent_rgb[0]:02X}{accent_rgb[1]:02X}{accent_rgb[2]:02X}")
    pBdr.append(bottom); pPr.append(pBdr)

def _add_section_heading_docx(doc: Document, label: str, accent_rgb: tuple):
    h = doc.add_paragraph(); h.paragraph_format.space_before = Pt(10); h.paragraph_format.space_after = Pt(4)
    r = h.add_run(label.upper()); r.font.size = Pt(11); r.font.bold = True; r.font.name = "Calibri"
    r.font.color.rgb = RGBColor(*accent_rgb)


# ══════════════════════════════════════════════════════════════════════════════
# PDF RENDERING
# ══════════════════════════════════════════════════════════════════════════════

def _pdf_styles(accent_hex: str, header_style: str = "centered") -> dict:
    styles = getSampleStyleSheet()
    ah = f"#{accent_hex}"
    tc = f"#{TEXT_HEX}"
    mc = f"#{MUTED_HEX}"
    ha = TA_CENTER if header_style == "centered" else TA_LEFT
    return {
        "name":     ParagraphStyle("N", parent=styles["Title"], fontSize=22,
                                   textColor=colors.HexColor(tc), alignment=ha,
                                   spaceAfter=2, fontName="Helvetica-Bold"),
        "title":    ParagraphStyle("T", parent=styles["Normal"], fontSize=12,
                                   textColor=colors.HexColor(ah), alignment=ha,
                                   spaceAfter=4, fontName="Helvetica-Bold"),
        "contact":  ParagraphStyle("C", parent=styles["Normal"], fontSize=9.5,
                                   textColor=colors.HexColor(mc), alignment=ha, spaceAfter=12),
        "heading":  ParagraphStyle("H", parent=styles["Heading2"], fontSize=11,
                                   textColor=colors.HexColor(ah), spaceBefore=10,
                                   spaceAfter=4, fontName="Helvetica-Bold"),
        "body":     ParagraphStyle("B", parent=styles["Normal"], fontSize=10.5,
                                   textColor=colors.HexColor(tc), spaceAfter=8, leading=14),
        "bullet":   ParagraphStyle("Bu", parent=styles["Normal"], fontSize=10.5,
                                   textColor=colors.HexColor(tc), spaceAfter=3,
                                   leftIndent=14, bulletIndent=2, leading=14),
        "sidebar_heading": ParagraphStyle("SH", parent=styles["Normal"], fontSize=9,
                                          textColor=colors.HexColor(ah), spaceBefore=8,
                                          spaceAfter=2, fontName="Helvetica-Bold"),
        "sidebar_body": ParagraphStyle("SB", parent=styles["Normal"], fontSize=9,
                                       textColor=colors.HexColor(tc), spaceAfter=2, leading=12),
        "letter_body": ParagraphStyle("LB", parent=styles["Normal"], fontSize=11,
                                      textColor=colors.HexColor(tc), spaceAfter=12, leading=16),
        "small":    ParagraphStyle("Sm", parent=styles["Normal"], fontSize=10,
                                   textColor=colors.HexColor(tc), spaceAfter=2),
    }


def render_resume_pdf(text: str, profile: dict, layout: Optional[dict] = None) -> bytes:
    sections     = _parse_resume_sections(text)
    order        = _section_order(layout)
    accent_h     = _accent_hex(layout)
    header_style = (layout or {}).get("header_style", "centered")
    uses_dividers = (layout or {}).get("uses_dividers", True)
    is_two_col   = (layout or {}).get("columns", 1) == 2 or (layout or {}).get("has_sidebar", False)
    styles       = _pdf_styles(accent_h, header_style)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter,
                            topMargin=0.6*inch, bottomMargin=0.6*inch,
                            leftMargin=0.75*inch, rightMargin=0.75*inch)
    story = []

    name  = profile.get("name") or "[Your Full Name]"
    title = profile.get("title") or ""
    contact_parts = [
        profile.get("email") or "[Your Email Address]",
        profile.get("phone") or "[Your Phone Number]",
        profile.get("location") or "[Your City, Country]",
        profile.get("linkedin_url") or "[Your LinkedIn URL]",
    ]
    if profile.get("github_url"): contact_parts.append(profile["github_url"])
    contact_text = " | ".join(contact_parts)

    # Banner header
    if header_style == "banner":
        ah_color = colors.HexColor(f"#{accent_h}")
        banner_data = [[Paragraph(f'<font size="18" color="white"><b>{_esc(name)}</b></font>' +
                                  (f'<br/><font size="10" color="#DDDDFF">{_esc(title)}</font>' if title else ""),
                                  ParagraphStyle("BH", alignment=TA_CENTER))]]
        banner_tbl = Table(banner_data, colWidths=[7*inch])
        banner_tbl.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,-1), ah_color),
            ("TOPPADDING",    (0,0), (-1,-1), 12),
            ("BOTTOMPADDING", (0,0), (-1,-1), 12),
            ("LEFTPADDING",   (0,0), (-1,-1), 12),
        ]))
        story.append(banner_tbl)
        story.append(Spacer(1, 6))
        story.append(Paragraph(_esc(contact_text), styles["contact"]))
    else:
        story.append(Paragraph(_esc(name), styles["name"]))
        if title: story.append(Paragraph(_esc(title), styles["title"]))
        story.append(Paragraph(_esc(contact_text), styles["contact"]))

    if uses_dividers:
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor(f"#{accent_h}")))
        story.append(Spacer(1, 6))

    if is_two_col:
        _pdf_two_col_body(story, sections, order, styles, accent_h)
    else:
        _pdf_single_col_body(story, sections, order, styles)

    doc.build(story)
    return buf.getvalue()


def _pdf_single_col_body(story: list, sections: dict, order: list, styles: dict):
    labels = {"SUMMARY":"Summary","SKILLS":"Skills","EXPERIENCE":"Experience",
              "PROJECTS":"Projects","EDUCATION":"Education","CERTIFICATIONS":"Certifications",
              "AWARDS":"Awards","LANGUAGES":"Languages","INTERESTS":"Interests",
              "OBJECTIVE":"Objective","PUBLICATIONS":"Publications"}
    for key in order:
        body = sections.get(key)
        if not body: continue
        story.append(Paragraph(labels.get(key,key.title()).upper(), styles["heading"]))
        if key in ("EXPERIENCE","PROJECTS"):
            for line in _bullets(body):
                story.append(Paragraph(f"&bull;&nbsp;&nbsp;{_esc(line)}", styles["bullet"]))
        else:
            story.append(Paragraph(_esc(body).replace("\n","<br/>"), styles["body"]))


def _pdf_two_col_body(story: list, sections: dict, order: list, styles: dict, accent_h: str):
    sidebar_keys = [k for k in ["SKILLS","CERTIFICATIONS","LANGUAGES","INTERESTS"] if sections.get(k)]
    main_keys    = [k for k in order if k not in sidebar_keys and k != "CONTACT" and sections.get(k)]

    if not sidebar_keys:
        _pdf_single_col_body(story, sections, order, styles)
        return

    # Build sidebar cell content
    left_content = []
    for key in sidebar_keys:
        left_content.append(Paragraph(key.title().upper(), styles["sidebar_heading"]))
        for item in _bullets(sections[key]):
            left_content.append(Paragraph(_esc(item), styles["sidebar_body"]))
        left_content.append(Spacer(1, 6))

    # Build main cell content
    right_content = []
    labels = {"SUMMARY":"Summary","EXPERIENCE":"Experience","PROJECTS":"Projects",
              "EDUCATION":"Education","OBJECTIVE":"Objective","AWARDS":"Awards","PUBLICATIONS":"Publications"}
    for key in main_keys:
        right_content.append(Paragraph(labels.get(key,key.title()).upper(), styles["heading"]))
        body = sections[key]
        if key in ("EXPERIENCE","PROJECTS"):
            for line in _bullets(body):
                right_content.append(Paragraph(f"&bull;&nbsp;{_esc(line)}", styles["bullet"]))
        else:
            right_content.append(Paragraph(_esc(body).replace("\n","<br/>"), styles["body"]))

    tbl = Table([[left_content, right_content]], colWidths=[2.0*inch, 5.0*inch])
    tbl.setStyle(TableStyle([
        ("VALIGN",        (0,0), (-1,-1), "TOP"),
        ("BACKGROUND",    (0,0), (0,-1),  colors.HexColor(f"#F4F3FF")),
        ("LEFTPADDING",   (0,0), (0,-1),  8),
        ("RIGHTPADDING",  (0,0), (0,-1),  8),
        ("LEFTPADDING",   (1,0), (1,-1),  12),
        ("LINEAFTER",     (0,0), (0,-1),  0.5, colors.HexColor(f"#{accent_h}")),
    ]))
    story.append(tbl)


def render_cover_letter_pdf(text: str, profile: dict, job: Optional[dict] = None) -> bytes:
    styles = _pdf_styles(DEFAULT_BRAND)
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter,
                            topMargin=inch, bottomMargin=inch,
                            leftMargin=inch, rightMargin=inch)
    story = []
    name  = profile.get("name") or "[Your Full Name]"
    today = datetime.now().strftime("%B %d, %Y")
    for line in [name, profile.get("email") or "[Your Email Address]",
                 profile.get("phone") or "[Your Phone Number]",
                 profile.get("location") or "[Your City, Country]"]:
        story.append(Paragraph(_esc(line), styles["small"]))
    story.append(Spacer(1, 14))
    story.append(Paragraph(_esc(today), styles["small"]))
    story.append(Spacer(1, 14))
    if job:
        if job.get("company"):
            story.append(Paragraph(_esc(job["company"]), styles["small"]))
        story.append(Paragraph("[Company Address]", styles["small"]))
        story.append(Spacer(1, 14))
    for para in [p.strip() for p in re.split(r"\n\s*\n", text.strip()) if p.strip()]:
        para = re.sub(r"(?i)\b(best regards|sincerely|kind regards|warm regards),\s*\n?\s*",
                      r"\1,<br/>", para)
        story.append(Paragraph(_esc_keep_br(para), styles["letter_body"]))
    doc.build(story)
    return buf.getvalue()


def _esc(text: str) -> str:
    return (text or "").replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")

def _esc_keep_br(text: str) -> str:
    placeholder = "\x00BR\x00"
    text = text.replace("<br/>", placeholder)
    text = _esc(text)
    return text.replace(placeholder, "<br/>")
