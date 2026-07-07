"""
Resume file parser — handles PDF and DOCX extraction, with an automatic
Gemini Vision OCR fallback for scanned/image-based PDFs.

Why this matters: PyPDF2 only reads embedded text layers. A scanned resume
(photographed or exported as flattened images) has NO text layer, so
PyPDF2 silently returns an empty/near-empty string — which previously
caused the app to reject the upload outright. This module now detects that
case and renders the PDF pages as images, then asks Gemini's multimodal
model to transcribe them, the same way a human would read a scanned page.
"""

from __future__ import annotations
import io
import os
import asyncio
from pathlib import Path

# Heuristic: if extracted text has fewer than this many non-whitespace
# characters per page, treat the PDF as "scanned" / image-based and trigger
# the OCR fallback rather than trusting the (near-empty) text layer.
MIN_CHARS_PER_PAGE = 40


def extract_text(file_bytes: bytes, filename: str) -> dict:
    """
    Extract plain text from a PDF or DOCX file.

    Returns a dict: {"text": str, "method": "text" | "ocr_vision" | "empty"}
    so callers (and the profile, in turn the UI) can be transparent about
    how the content was obtained — important since OCR is a best-effort
    transcription and may contain minor errors.
    """
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return _extract_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return {"text": _extract_docx(file_bytes), "method": "text"}
    elif ext == ".txt":
        return {"text": file_bytes.decode("utf-8", errors="replace"), "method": "text"}
    else:
        try:
            return {"text": file_bytes.decode("utf-8", errors="replace"), "method": "text"}
        except Exception:
            return {"text": "", "method": "empty"}


def _extract_pdf(file_bytes: bytes) -> dict:
    """Try the fast text-layer extraction first; fall back to OCR if it looks scanned."""
    text, page_count = _extract_pdf_text_layer(file_bytes)

    looks_scanned = page_count > 0 and len(text.strip()) < (MIN_CHARS_PER_PAGE * page_count)

    if not looks_scanned and text.strip():
        return {"text": text, "method": "text"}

    # Text layer is empty or suspiciously thin → likely a scanned/image PDF.
    # Attempt Gemini Vision OCR as a fallback so scanned resumes still work.
    ocr_text = _extract_pdf_via_vision_ocr(file_bytes)
    if ocr_text.strip():
        return {"text": ocr_text, "method": "ocr_vision"}

    # Neither path produced usable text — return whatever we have (possibly
    # empty); the caller decides how to handle a fully empty result.
    return {"text": text, "method": "empty" if not text.strip() else "text"}


def _extract_pdf_text_layer(file_bytes: bytes) -> tuple[str, int]:
    """Fast path: read the embedded text layer via PyPDF2."""
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        pages = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                pages.append(t)
        return "\n".join(pages), len(reader.pages)
    except Exception:
        return "", 0


def _extract_pdf_via_vision_ocr(file_bytes: bytes, max_pages: int = 5) -> str:
    """
    Render PDF pages to images with PyMuPDF and transcribe them using
    Gemini's multimodal vision capability. Capped at `max_pages` to keep
    latency and token usage bounded for the free tier — a resume is
    virtually never longer than this.
    """
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        return ""

    try:
        import fitz  # PyMuPDF
        import google.generativeai as genai

        genai.configure(api_key=api_key)

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        images = []
        for i, page in enumerate(doc):
            if i >= max_pages:
                break
            # 2x zoom for sharper OCR on small/dense resume fonts
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            images.append({"mime_type": "image/png", "data": pix.tobytes("png")})
        doc.close()

        if not images:
            return ""

        model = genai.GenerativeModel(os.getenv("GEMINI_MODEL", "gemini-2.5-flash"))
        prompt = (
            "This is a scanned resume/CV. Transcribe ALL visible text exactly as it "
            "appears, preserving section structure (name, contact info, summary, "
            "experience, education, skills, etc.) using plain text with line breaks "
            "between sections. Do not summarise or omit anything — produce a complete, "
            "faithful transcription so it can be parsed programmatically afterwards."
        )

        parts = [prompt] + images
        response = model.generate_content(parts)
        return response.text or ""
    except Exception:
        return ""


def _extract_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception as e:
        return f"[DOCX extraction failed: {e}]"
